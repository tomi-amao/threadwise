"""
Embedding Service for AI Agent
Handles text extraction from various file types and embedding generation
"""

import os
import tempfile
from io import BytesIO
from typing import List, Dict, Any, Optional, Tuple
import logging
from pathlib import Path
import asyncio


import httpx
try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from sentence_transformers import SentenceTransformer
    from langchain_community.embeddings import SentenceTransformerEmbeddings
    from langchain_huggingface import HuggingFaceEmbeddings
except ImportError:
    SentenceTransformer = None
    SentenceTransformerEmbeddings = None

from langchain_community.vectorstores import SupabaseVectorStore
from langchain_core.documents import Document as LangChainDocument
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
try:
    from supabase import create_client, Client
except ImportError:
    create_client = None
    Client = None

logger = logging.getLogger(__name__)

# Configuration
SUPABASE_URL = os.getenv("SUPABASE_URL", "http://localhost:8000")
SUPABASE_SERVICE_ROLE_KEY = os.getenv("SUPABASE_SERVICE_ROLE_KEY", "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyAgCiAgICAicm9sZSI6ICJzZXJ2aWNlX3JvbGUiLAogICAgImlzcyI6ICJzdXBhYmFzZS1kZW1vIiwKICAgICJpYXQiOiAxNjQxNzY5MjAwLAogICAgImV4cCI6IDE3OTk1MzU2MDAKfQ.DaYlNEoUrrEn2Ig7tqibS-PHK5vgusbcbo7X36XVt4Q")

# Embedding model configuration
EMBEDDING_MODEL = "BAAI/bge-small-en"
VECTOR_TABLE_NAME = "documents"
VECTOR_QUERY_NAME = "match_documents"

# Global instances (initialized once)
_embedding_model = None
_embeddings = None
_supabase_client = None





class EmbeddingService:
    """Service for handling file embedding operations"""
    
    def __init__(self):
        self.embeddings = self._get_embeddings()
        self.supabase = self._get_supabase_client()
    
    def _get_embeddings(self):
        """Get or initialize the embedding model"""
        global _embeddings
        if _embeddings is None:
            if SentenceTransformerEmbeddings is None:
                raise ImportError("sentence-transformers not available")
            logger.info(f"Initializing embedding model: {EMBEDDING_MODEL}")
            _embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL)
            logger.info(f"Embedding model initialized successfully {_embeddings}")
        return _embeddings
    
    def _get_supabase_client(self):
        """Get or initialize the Supabase client"""
        global _supabase_client
        if _supabase_client is None:
            if create_client is None:
                raise ImportError("supabase not available")
            logger.info(f"Initializing Supabase client: {SUPABASE_URL}")
            _supabase_client = create_client(SUPABASE_URL, SUPABASE_SERVICE_ROLE_KEY)
            logger.info("Supabase client initialized successfully")
        return _supabase_client
    
    async def download_file_from_url(self, file_url: str) -> bytes:
        """Download file content from a URL"""
        async with httpx.AsyncClient() as client:
            response = await client.get(file_url)
            response.raise_for_status()
            return response.content
    
    def extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            reader = PdfReader(BytesIO(file_content))
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            raise ValueError(f"Failed to extract text from PDF: {e}")
    
    def extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from Word document"""
        try:
            if DocxDocument is None:
                raise ImportError("python-docx not available")
            doc = DocxDocument(BytesIO(file_content))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            raise ValueError(f"Failed to extract text from DOCX: {e}")
    
    def extract_text_from_txt(self, file_content: bytes) -> str:
        """Extract text from plain text file"""
        try:
            # Try different encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            raise ValueError("Could not decode text file with any common encoding")
        except Exception as e:
            logger.error(f"Error extracting text from TXT: {e}")
            raise ValueError(f"Failed to extract text from TXT: {e}")
    
    def extract_text_from_file(self, file_content: bytes, file_type: str, filename: str) -> str:
        """Extract text from file based on file type"""
        logger.info(f"Extracting text from {filename} (type: {file_type})")
        
        file_type = file_type.lower()
        
        if file_type == 'application/pdf':
            return self.extract_text_from_pdf(file_content)
        elif file_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
            return self.extract_text_from_docx(file_content)
        elif file_type.startswith('text/') or file_type in ['application/json', 'text/plain']:
            return self.extract_text_from_txt(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """Chunk text into smaller pieces for better embedding"""
        if len(text) <= chunk_size:
            return [text] if text.strip() else []
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = min(start + chunk_size, len(text))
            chunk = text[start:end]
            
            # Try to break at sentence or paragraph boundaries
            if end < len(text):
                # Look for sentence endings
                last_sentence = max(
                    chunk.rfind('.'),
                    chunk.rfind('!'),
                    chunk.rfind('?'),
                    chunk.rfind('\n\n')
                )
                
                if last_sentence > start + chunk_size * 0.5:
                    chunk = text[start:last_sentence + 1]
                    start = last_sentence + 1 - overlap
                else:
                    start = end - overlap
            else:
                start = end
            
            if chunk.strip():
                chunks.append(chunk.strip())
        
        logger.info(f"Text chunked into {len(chunks)} pieces")
        return chunks
    async def embed_file(self, file_url: str, filename: str, file_type: str) -> Dict[str, Any]:
        """
        Process a file: download, extract text, create embeddings, and store in vector database
        """
        logger.info(f"Starting embedding process for: {filename}")
        
        try:
            # Create and load PDF documents in a thread to avoid blocking
            def load_and_split_pdf():
                # Load PDF
                pdfLoader = PyPDFLoader(
                    file_path=file_url,
                    mode="single",
                    pages_delimiter="\n"
                )
                documents = pdfLoader.load()  # Use sync load() in the thread
                
                # Split documents
                splitter = RecursiveCharacterTextSplitter(
                    chunk_size=500,
                    chunk_overlap=50,
                    separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""]
                )
                return splitter.split_documents(documents)

            documents = await asyncio.to_thread(load_and_split_pdf)
            print(documents[0].page_content[:100])
            print(documents[0].metadata)

            # Store in vector database
            logger.info(f"Storing {len(documents)} document chunks in vector database")
            
            # Wrap the vector store creation in asyncio.to_thread
            def create_vector_store():
                return SupabaseVectorStore.from_documents(
                    documents,
                    self.embeddings,
                    client=self.supabase,
                    table_name=VECTOR_TABLE_NAME,
                    query_name=VECTOR_QUERY_NAME,
                )
            
            vector_store = await asyncio.to_thread(create_vector_store)
            
            logger.info(f"Successfully embedded {len(documents)} chunks for {filename}")
            
            return {
                "success": True,
                "documentId": documents[0].metadata.get("documentId", f"{filename}-{hash(file_url)}"),
                "chunks": len(documents),
                "filename": filename,
                "text_length": sum(len(doc.page_content) for doc in documents),
            }
            
        except Exception as e:
            logger.error(f"Error embedding file {filename}: {e}")
            return {
                "success": False,
                "error": str(e),
                "filename": filename,
            }
    
    async def search_documents(self, query: str, limit: int = 5) -> List[Dict[str, Any]]:
        """Search for similar documents"""
        try:
            logger.info(f"Searching documents for: {query}")
            
            vector_store = SupabaseVectorStore(
                embedding=self.embeddings,
                client=self.supabase,
                table_name=VECTOR_TABLE_NAME,
                query_name=VECTOR_QUERY_NAME,
            )


            async def async_rpc(client, query_embedding, match_count=4):
                return await asyncio.to_thread(
                    lambda: client.rpc(VECTOR_QUERY_NAME, 
    {
                        "query_embedding": query_embedding,
                        "match_count": match_count,
                    }).execute()
                )
            query_embedding = vector_store.embeddings.embed_query(query)
            similarity_search_response = await async_rpc(vector_store._client, query_embedding, 4)

            results = [
                LangChainDocument(
                    page_content=row['content'],
                    metadata=row['metadata']
                )
                for row in similarity_search_response.data
            ]

            # results = await vector_store.asearch(query=query, search_type="similarity")
            logger.info(f"Search returned {len(results)} results")
            logger.info(f"Search results: {results}")
            
            search_results = []
            for doc in results:
                search_results.append({
                    "content": doc.page_content,
                    "metadata": {
                        # "fileName": doc.metadata.get("fileName"),
                        # "fileType": doc.metadata.get("fileType"),
                        # "chunkIndex": doc.metadata.get("chunkIndex"),
                        # "score": float(score),
                        "source": doc.metadata.get("source"),
                    }
                })
            
            # logger.info(f"Found {len(search_results)} similar documents")
            return search_results

        except Exception as e:
            logger.error(f"Error searching documents: {e}")
            return [{"error": str(e)}]

    async def delete_file_embeddings(self, filename: str) -> bool:
        """Delete all embeddings for a specific file"""
        try:
            logger.info(f"Deleting embeddings for: {filename}")
            
            # Delete from the documents table where metadata contains the filename
            # Wrap in asyncio.to_thread to avoid blocking operations
            def delete_embeddings():
                return self.supabase.table(VECTOR_TABLE_NAME).delete().eq('metadata->>fileName', filename).execute()
            
            result = await asyncio.to_thread(delete_embeddings)
            
            if result.data is not None:
                logger.info(f"Successfully deleted embeddings for {filename}")
                return True
            else:
                logger.warning(f"No embeddings found for {filename}")
                return False
                
        except Exception as e:
            logger.error(f"Error deleting embeddings for {filename}: {e}")
            return False
    
    async def health_check(self) -> Dict[str, Any]:
        """Check if the embedding service is healthy"""
        try:
            # Test embedding model - wrap in asyncio.to_thread to avoid blocking
            def test_embedding():
                test_text = "This is a test."
                return self.embeddings.embed_query(test_text)
            
            test_embedding = await asyncio.to_thread(test_embedding)
            
            # Test Supabase connection - wrap in asyncio.to_thread to avoid blocking
            def test_supabase():
                return self.supabase.table(VECTOR_TABLE_NAME).select("id").limit(1).execute()
            
            result = await asyncio.to_thread(test_supabase)
            
            return {
                "status": "healthy",
                "embedding_model": EMBEDDING_MODEL,
                "embedding_dimensions": len(test_embedding),
                "vector_store": "connected",
            }
        except Exception as e:
            logger.error(f"Health check failed: {e}")
            return {
                "status": "unhealthy",
                "error": str(e),
            }


# Global service instance
embedding_service = EmbeddingService()
